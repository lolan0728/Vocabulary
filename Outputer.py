###############################################################################
# Lolan 2020/08/06
###############################################################################

import datetime
import os
from openpyxl.descriptors.base import Length
from xlsxwriter import Workbook, worksheet
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt
from Vocabulary import WordType

class Outputer:
    
    def __init__(self):
        pass
    
    def outputDocx(self, title, name, savepath, 
                        wordsTable, date=None):
        
        date = datetime.date.today().strftime("%Y/%m/%d") \
                if date is None else date

        new_words, review_Words = self.formatWordsTable(wordsTable)

        new_words = self._split_2(new_words, 4)
        review_Words = self._split_2(review_Words, 14)
        
        document = Document()
        document.styles['Normal'].font.name = u'Times New Roman'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        section = document.sections[0]
        section.top_margin=Cm(1)
        section.bottom_margin=Cm(1)
        section.left_margin=Cm(2)
        section.right_margin=Cm(2)

        title_ = document.add_heading(level=0)
        # title_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_.add_run(title)
        title_run.font.size = Pt(24)
        # title_run.font.name = 'Times New Roman'
        # title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title_run.bold = True

        date_ = document.add_paragraph()
        date_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_.add_run(date)
        date_run.font.size = Pt(14)
        date_run.font.name = 'Times New Roman'
        # date_run.italic = True

        new_ = document.add_heading(level=1)
        new_run = new_.add_run('New Words')
        new_run.font.size = Pt(18)

        table_new = document.add_table(rows=0, cols=2)
        table_new.autofit=True
        for a, b in new_words:
            row = table_new.add_row()
            row_cells = row.cells
            row_cells[0].text = a
            row_cells[1].text = b 
        
        review_ = document.add_heading(level=1)
        review_run = review_.add_run('Review')
        review_run.font.size = Pt(18)

        table_review = document.add_table(rows=0, cols=2)
        table_review.autofit=True
        for a, b in review_Words:
            row = table_review.add_row()
            row_cells = row.cells
            row_cells[0].text = a
            row_cells[1].text = b 

        self._setTableFontSize(table_new)
        self._setTableFontSize(table_review)

        # path = os.path.join(savepath, filename)
        document.save(savepath)

    def _setTableFontSize(self, table, fontsize=14):
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        # run.font.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                        font = run.font
                        font.size= Pt(fontsize)  
    
    def formatWordsTable(self, wordsTable):
        new_words = []
        review_Words = []
        for w in wordsTable.values():
            # mark = '★' if w.type == WordType.NEW else '   '
            # word = '    %s: %s' % (w.word, w.explanation)
            word = '    %s      %s' % (w.word, w.explanation)
            # word = '□ □ □ %s: %s' % (w.word, w.explanation)
            if w.type == WordType.NEW:
                new_words.append(word)
            else:
                review_Words.append(word)
        return new_words, review_Words

    def _split_2(self, records, rowNum):
        result = []
        while len(records) > 0:
            blockSize = rowNum*2 if len(records) > rowNum*2-1 else len(records)
            ls = records[:blockSize]
            records = records[blockSize:]
            if len(ls) > rowNum:
                part1 = ls[:rowNum]
                part2 = ls[rowNum:] + ['' for i in range(rowNum*2-len(ls))]
            else:
                part1 = ls
                part2 = ['' for i in range(len(ls))]
            result = [w for w in list(zip(part1, part2))]
        return result

    def outputXlsx(self, outVocabs, savepath):
        wb = Workbook(savepath)
        
        titlefomat = wb.add_format({
            'font_size': 16,
            'bold':  True,                 
            'border':1,                    
            'valign':   'vcenter',         
            'fg_color': '#02C874'
        })

        datefomat = wb.add_format({
            'font_size': 7,                
            'align':    'center',          
            'valign':   'vcenter',
            'border':0,                    
        })

        wordfomat = wb.add_format({
            'font_size': 13,               
            'valign':   'vcenter',
            'border':1,    
        })

        wordNotYetfomat = wb.add_format({
            'font_size': 13,               
            'valign':   'vcenter',
            'border':1,           
            'fg_color': '#DCDCDC'         
        })

        expfomat = wb.add_format({
            'font_size': 11,               
            'valign':   'vcenter',         
            'border':1,                    
        })

        newfomat = wb.add_format({
            'fg_color': 'red',
            'border':1,                    
        })

        reviewfomat = wb.add_format({
            'fg_color': 'yellow',
            'border':1,                    
        })

        nonefomat = wb.add_format({
            'border':1,                    
        })
        for vocab in outVocabs:
            ws = wb.add_worksheet(vocab.name)

            # word column width
            ws.set_column('A:A', 16) 
            # explanation column width
            ws.set_column('B:B', 25) 
            # date column width
            if len(vocab.alldates) > 0:
                ws.set_column(2, len(vocab.alldates)+1, 5) 

            # title
            ws.merge_range('A1:B2', vocab.description, titlefomat)

            # date row
            row = 1
            col = 2
            for d in vocab.alldates:
                # date = datetime.datetime.strptime(d, '%Y/%m/%d').strftime("%m/%d")
                date = datetime.datetime.strptime(d, '%Y/%m/%d').strftime("%y/%m/%d")
                ws.write(row, col, date, datefomat)            
                col += 1
            row += 1

            col = 2
            # words
            sortedWords = list(vocab.words.keys())
            # sortedWords.sort()
            for w in sortedWords:
                word = vocab.words[w]
                # word
                ws.write(row, 0, word.word, wordfomat if(word.newDate) else wordNotYetfomat)

                # explanation
                ws.write(row, 1, word.explanation, expfomat)
                # date mark
                for i in range(len(vocab.alldates)):
                    date = vocab.alldates[i]
                    if word.newDate == date:
                        ws.write(row, col+i, '', newfomat)
                    elif date in word.reviewDate:
                        ws.write(row, col+i, '', reviewfomat)
                    else:
                        ws.write(row, col+i, '', nonefomat)
                ws.set_row(row, 28)
                row += 1
            ws.freeze_panes(2, 2)
        wb.close()

    def outputCharts(self, mPData, vData, sData, savepath):
        wb = Workbook(savepath)

        # Monthly performance
        boldTitle = wb.add_format({'bold': 1})
        mpDataS = wb.add_worksheet('MPDataSource')
        vDataS = wb.add_worksheet('VDataSource')
        sDataS = wb.add_worksheet('SDataSource')

        self.writeDataSource(mpDataS, 0, boldTitle, mPData)
        self.writeDataSource(vDataS, 0, boldTitle, vData)
        self.writeDataSource(sDataS, 0, boldTitle, sData)

        s_pos = [0, 0]
        s_dataSheetName = 'SDataSource'
        ws_s = wb.add_worksheet('Summary')
        self.drawDoughnutChart(wb, ws_s, s_pos, s_dataSheetName, 2, [500, 400])
        s_pos[0] += 23
        r = 0
        for i in range(3, len(sData)+1):
            self.drawDoughnutChart(wb, ws_s, s_pos, s_dataSheetName, i, [350, 300])
            s_pos[1] += 6
            r += 1
            if r > 3:
                s_pos[0] += 17
                s_pos[1] = 0
                r = 0

        mp_chartTitle = '月表现'
        mp_xAxisName = 'Year/Month'
        # mp_yAxisName = 'Number of words'
        mp_yAxisName = ''
        mp_pos = [0, 0]
        mp_dataSheetName = 'MPDataSource'
        ws_h = wb.add_worksheet('Monthly Performance')
        self.drawHistogram(wb, ws_h, mp_chartTitle, mp_xAxisName, 
                                mp_yAxisName, mp_pos, mp_dataSheetName, len(mPData))
        mpDataS.hide()

        v_chartTitle = '词汇量'
        v_xAxisName = 'Year/Month'
        v_yAxisName = 'Number of words'
        v_pos = [0, 0]
        v_dataSheetName = 'VDataSource'
        ws_v = wb.add_worksheet('Trend')
        self.drawLineChart(wb, ws_v, v_chartTitle, v_xAxisName, v_yAxisName, v_pos, v_dataSheetName, len(vData))
        vDataS.hide()
                
        ws_s.activate()

        sDataS.hide()
        wb.close()

    def writeDataSource(self, ws, strRow, titleFormat, data):
        title = data[0]
        data = data[1:]

        ws.write_row(strRow, 0, title, titleFormat)
        strRow += 1
        for line in data:
            ws.write_row(strRow, 0, line)
            strRow += 1
        

    def drawHistogram(self, wb, ws, chartTitle, xAxisName, yAxisName, position, dataSheetName, datalines):
        # ws = wb.add_worksheet(sheetName)
        chart_col = wb.add_chart({'type': 'column'})

        series_name_new = '=' + dataSheetName + '!$B$1'
        series_name_review = '=' + dataSheetName + '!$C$1'
        series_name_workdays = '=' + dataSheetName + '!$D$1'
        series_cat = '=' + dataSheetName + '!$A$2:$A$' + str(datalines)
        series_values_new = '=' + dataSheetName + '!$B$2:$B$' + str(datalines)
        series_values_review = '=' + dataSheetName + '!$C$2:$C$' + str(datalines)
        series_values_workdays = '=' + dataSheetName + '!$D$2:$D$' + str(datalines)

        chart_col.add_series({
            'name': series_name_new,
            'categories': series_cat,
            'values':   series_values_new,
            'data_labels': {'value': True},
            # 'line': {'color': 'gray'},
        })

        chart_col.add_series({
            'name': series_name_review,
            'categories':  series_cat,
            'values':   series_values_review,
            'data_labels': {'value': True},
            # 'line': {'color': 'yellow'},
        })

        chart_col.add_series({
            'name': series_name_workdays,
            'categories':  series_cat,
            'values':   series_values_workdays,
            'data_labels': {'value': True},
            # 'line': {'color': 'yellow'},
        })

        chart_col.set_title({'name': chartTitle})
        chart_col.set_x_axis({'name': xAxisName})
        chart_col.set_y_axis({'name': yAxisName})

        chart_col.set_style(2)
        chart_col.height = 600
        chart_col.width = 1200

        ws.insert_chart(position[0], position[1], chart_col, {'x_offset': 25, 'y_offset': 10})
        # ws.activate()

    def drawLineChart(self, wb, ws, chartTitle, xAxisName, yAxisName, position, dataSheetName, datalines):
        # ws = wb.add_worksheet(sheetName)
        chart_col = wb.add_chart({'type': 'line'})

        series_name_new = '=' + dataSheetName + '!$B$1'
        series_cat = '=' + dataSheetName + '!$A$2:$A$' + str(datalines)
        series_values = '=' + dataSheetName + '!$B$2:$B$' + str(datalines)

        chart_col.add_series({
            'name': series_name_new,
            'categories': series_cat,
            'values':   series_values,
            'line': {'color': '#20B2AA'},
            'marker': {
                'type': 'square',
                'size': 8,
                'border': {'color': '#20B2AA'},
                'fill':   {'color': '#20B2AA'},
            },
            'data_labels': {'value': True},
            'smooth': True,
        })

        chart_col.set_title({'name': chartTitle})
        chart_col.set_x_axis({'name': xAxisName})
        chart_col.set_y_axis({'name': yAxisName})

        chart_col.set_style(12)
        chart_col.height = 600
        chart_col.width = 1200

        ws.insert_chart(position[0], position[1], chart_col, {'x_offset': 25, 'y_offset': 10})
        # ws.activate()

    def drawDoughnutChart(self, wb, ws, position, dataSheetName, datalines, size=[300,200]):
        # ws = wb.add_worksheet(sheetName)
        chart_col = wb.add_chart({'type': 'doughnut'})

        series_name = '=' + dataSheetName + '!$A$' + str(datalines)
        series_cat = '=' + dataSheetName + '!$B$1:$E$1'
        series_values = '=' + dataSheetName + '!$B$'+ str(datalines) + ':$E$' + str(datalines)

        chart_col.add_series({
            'name': series_name,
            'categories': series_cat,
            'values':   series_values,
            # 'line': {'color': 'red'},
            # 'data_labels': {'value': True},
            # 'data_labels': {'percentage': True, 'leader_lines': True},

            'points': [
                {'fill': {'color': '#D3D3D3'}},
                {'fill': {'color': '#00CED1'}},
                {'fill': {'color': '#32CD32'}},
                {'fill': {'color': '#DAA520'}},
                # {'fill': {'color': '#F4A460'}},
            ],
        })

        chart_col.set_title({'name': series_name})
        # chart_col.set_rotation(28)
        chart_col.set_style(9)
        chart_col.set_hole_size(60)
        chart_col.width = size[0]
        chart_col.height = size[1]

        ws.insert_chart(position[0], position[1], chart_col, {'x_offset': 25, 'y_offset': 10})
        # ws.activate()

    def tester(self, savepath):
        # 创建一个excel
        workbook = Workbook(savepath)
        # 创建一个sheet
        worksheet = workbook.add_worksheet()
        # worksheet = workbook.add_worksheet("bug_analysis")

        # 自定义样式，加粗
        bold = workbook.add_format({'bold': 1})

        # --------1、准备数据并写入excel---------------
        # 向excel中写入数据，建立图标时要用到
        headings = ['Number', 'testA', 'testB']
        data = [
            ['2017-9-1', '2017-9-2', '2017-9-3', '2017-9-4', '2017-9-5', '2017-9-6'],
            [10, 40, 50, 20, 10, 50],
            [30, 60, 70, 50, 40, 30],
        ]

        # 写入表头
        worksheet.write_row('A1', headings, bold)

        # 写入数据
        worksheet.write_column('A2', data[0])
        worksheet.write_column('B2', data[1])
        worksheet.write_column('C2', data[2])

        # --------2、生成图表并插入到excel---------------
        # 创建一个柱状图(column chart)
        chart_col = workbook.add_chart({'type': 'column'})

        # 配置第一个系列数据
        chart_col.add_series({
            # 这里的sheet1是默认的值，因为我们在新建sheet时没有指定sheet名
            # 如果我们新建sheet时设置了sheet名，这里就要设置成相应的值
            'name': '=Sheet1!$B$1',
            'categories': '=Sheet1!$A$2:$A$7',
            'values':   '=Sheet1!$B$2:$B$7',
            'line': {'color': 'red'},
        })

        # 配置第二个系列数据(用了另一种语法)
        chart_col.add_series({
            'name': '=Sheet1!$C$1',
            'categories':  '=Sheet1!$A$2:$A$7',
            'values':   '=Sheet1!$C$2:$C$7',
            'line': {'color': 'yellow'},
        })

        # 配置第二个系列数据(用了另一种语法)
        # chart_col.add_series({
        #     'name': ['Sheet1', 0, 2],
        #     'categories': ['Sheet1', 1, 0, 6, 0],
        #     'values': ['Sheet1', 1, 2, 6, 2],
        #     'line': {'color': 'yellow'},
        # })

        # 设置图表的title 和 x，y轴信息
        chart_col.set_title({'name': 'The xxx site Bug Analysis'})
        chart_col.set_x_axis({'name': 'Test number'})
        chart_col.set_y_axis({'name':  'Sample length (mm)'})

        # 设置图表的风格
        chart_col.set_style(2)

        # 把图表插入到worksheet以及偏移
        worksheet.insert_chart('A10', chart_col, {'x_offset': 25, 'y_offset': 10})

        workbook.close()


if __name__ == "__main__":
    savepath = r'C:\Users\vampi\Desktop\test.xlsx'
    Outputer().outputCharts(savepath)
    # Outputer().tester(savepath)
    #self tester
    # import Logic as logic
    # t = logic.Logic()
    # t.loadAll()
    # savepath = r'C:\Users\eos\Desktop\test.xlsx'
    # alldates = ['2020/08/01', '2020/9/10', '2020/9/11', '2020/9/12', '2020/9/13', '2020/9/14', '2020/9/15', '2020/12/31']
    # wordDate = {}
    # for w in t.vocabList['vocab_1'].words:
    #     word = t.uniVocabulary[w]
    #     if word.word == 'brother':
    #         word.newDate = '2020/08/01'
    #     if word.word == 'gun':
    #         word.newDate = '2020/12/31'
    #     if word.word == 'granddaughter':
    #         word.addReviewDate('2020/9/11')
    #         word.addReviewDate('2020/9/12')
    #         word.addReviewDate('2020/9/15')
    #     wordDate[w] = word
    
    # Outputer().outputXlsx('vocab_1', '小学必备分类英语单词800个', alldates, wordDate, savepath)
    # newW = (
    #         ('    English: 英国的，英语', ''ttt),
    #         ('    English: 英国的，英语', ''),
    #         ('    English: 英国的，英语', ''),
    #     )
    # reviewW = (
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #         ('    English: 英国的，英语', '    France: 法国的，法语'),
    #     )
    # print(records)
            
    # Outputer().outputDocx("Harvey's Daily Words", 
    #                         'Harvey', 
    #                         r'C:\\Users\\eos\\Desktop\\', 
    #                         'test.docx', 
    #                         newW, 
    #                         reviewW)